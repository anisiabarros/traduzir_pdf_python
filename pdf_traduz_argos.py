#!/usr/bin/env python3
import os
import sys
import argparse
from pdf2docx import Converter           # pip install pdf2docx
from docx import Document                # pip install python-docx
import argostranslate.package as pkg      # pip install argostranslate
from argostranslate.translate import get_installed_languages, get_language_from_code
from tqdm import tqdm                    # pip install tqdm

def ensure_argos_model_installed(src_code, dest_code):
    # Atualiza índice e busca pacotes disponíveis
    pkg.update_package_index()
    available_packages = pkg.get_available_packages()
    # Encontra pacote de src->dest
    model_pkg = next(
        (p for p in available_packages if p.from_code == src_code and p.to_code == dest_code),
        None
    )
    if model_pkg is None:
        print(f"❌ Nenhum modelo Argos disponível para {src_code}->{dest_code}", file=sys.stderr)
        sys.exit(1)
    # Faz download e instala
    print("Baixando modelo Argos...", end="", flush=True)
    pkg_path = model_pkg.download()
    print(" concluído.")
    print("Instalando modelo Argos...", end="", flush=True)
    pkg.install_from_path(pkg_path)
    print(" concluído.")

def get_offline_translation_func(src_code, dest_code):
    # Carrega linguagens instaladas
    langs = get_installed_languages()
    from_lang = get_language_from_code(src_code)
    to_lang   = get_language_from_code(dest_code)
    translation = from_lang.get_translation(to_lang)
    def translate_text_block(text):
        paragraphs = text.split("\n\n")
        out = []
        for p in paragraphs:
            if not p.strip():
                out.append("")
            else:
                hypos = translation.hypotheses(p, num_hypotheses=1)
                out.append(hypos[0].value if hypos else p)
        return "\n\n".join(out)
    return translate_text_block

def pdf_to_docx(input_pdf, output_docx):
    converter = Converter(input_pdf)
    converter.convert(output_docx, start=0, end=None)
    converter.close()

def translate_docx(input_docx, output_docx, translator):
    doc = Document(input_docx)
    # Traduz parágrafos
    paras = [para for para in doc.paragraphs if para.text.strip()]
    for para in tqdm(paras, desc="Traduzindo parágrafos"):
        orig = para.text
        tr = translator(orig)
        if tr != orig:
            para.clear()
            para.add_run(tr)
    # Traduz tabelas
    cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cells.append(cell)
    for cell in tqdm(cells, desc="Traduzindo células de tabela"):
        orig = cell.text
        tr = translator(orig)
        if tr != orig:
            cell.text = tr
    doc.save(output_docx)

def main():
    parser = argparse.ArgumentParser(description="PDF→DOCX→tradução offline Argos→DOCX com barra de progresso")
    parser.add_argument("input_pdf", help="PDF original")
    parser.add_argument("intermediate_docx", help="DOCX gerado")
    parser.add_argument("output_docx", help="DOCX traduzido")
    parser.add_argument("--src", default="en", help="código ISO do idioma de origem")
    parser.add_argument("--dest", default="pt", help="código ISO do idioma de destino")
    args = parser.parse_args()

    if not os.path.isfile(args.input_pdf):
        print("Erro: PDF não encontrado:", args.input_pdf, file=sys.stderr)
        sys.exit(1)

    print(f"Instalando modelo Argos {args.src}->{args.dest} (se necessário)…")
    ensure_argos_model_installed(args.src, args.dest)

    print("Convertendo PDF para DOCX…")
    pdf_to_docx(args.input_pdf, args.intermediate_docx)
    print("Concluído conversão.")

    print("Carregando tradutor offline…")
    translator = get_offline_translation_func(args.src, args.dest)

    print("Iniciando tradução do DOCX...")
    translate_docx(args.intermediate_docx, args.output_docx, translator)

    print("Concluído! Revise", args.output_docx, "e exporte para PDF.")

if __name__ == "__main__":
    main()
