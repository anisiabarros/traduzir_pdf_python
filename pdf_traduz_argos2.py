#!/usr/bin/env python3
import os
import sys
import argparse
from pdf2docx import Converter           # pip install pdf2docx
from docx import Document                # pip install python-docx
import argostranslate.package as pkg      # pip install argostranslate
from argostranslate.translate import get_installed_languages, get_language_from_code

def ensure_argos_model_installed(src_code, dest_code):
    # 1) Atualiza índice e busca pacotes disponíveis
    pkg.update_package_index()
    available_packages = pkg.get_available_packages()
    # 2) Encontra pacote de src->dest
    model_pkg = next(
        (p for p in available_packages if p.from_code == src_code and p.to_code == dest_code),
        None
    )
    if model_pkg is None:
        print(f"❌ Nenhum modelo Argos disponível para {src_code}->{dest_code}", file=sys.stderr)
        sys.exit(1)
    # 3) Faz download e instala
    pkg_path = model_pkg.download()
    pkg.install_from_path(pkg_path)

def get_offline_translation_func(src_code, dest_code):
    # Carrega linguagens instaladas após instalar pacote
    langs = get_installed_languages()
    from_lang = get_language_from_code(src_code)
    to_lang   = get_language_from_code(dest_code)
    translation = from_lang.get_translation(to_lang)
    def translate_text_block(text):
        # Divide em parágrafos para manter quebras
        paragraphs = text.split("\n\n")
        out = []
        for p in paragraphs:
            if not p.strip():
                out.append("")
            else:
                hypos = translation.hypotheses(p, num_hypotheses=1)
                out.append(hypos[0].value if hypos else p)
        return "\n\n".join(out)
    return translate_text_block

def pdf_to_docx(input_pdf, output_docx):
    cv = Converter(input_pdf)
    cv.convert(output_docx, start=0, end=None)
    cv.close()

def translate_docx(input_docx, output_docx, translator):
    doc = Document(input_docx)
    # Parágrafos
    for para in doc.paragraphs:
        orig = para.text
        tr = translator(orig)
        if tr != orig:
            para.clear()
            para.add_run(tr)
    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                orig = cell.text
                tr = translator(orig)
                if tr != orig:
                    cell.text = tr
    doc.save(output_docx)

def main():
    parser = argparse.ArgumentParser(description="PDF→DOCX→tradução offline Argos→DOCX")
    parser.add_argument("input_pdf", help="PDF original")
    parser.add_argument("intermediate_docx", help="DOCX gerado")
    parser.add_argument("output_docx", help="DOCX traduzido")
    parser.add_argument("--src", default="en", help="código ISO 639 do idioma de origem")
    parser.add_argument("--dest", default="pt", help="código ISO 639 do idioma de destino")
    args = parser.parse_args()

    if not os.path.isfile(args.input_pdf):
        print("Erro: PDF não encontrado:", args.input_pdf, file=sys.stderr)
        sys.exit(1)

    print(f"Instalando modelo Argos {args.src}->{args.dest} (se necessário)…")
    ensure_argos_model_installed(args.src, args.dest)

    print("Convertendo PDF para DOCX…")
    pdf_to_docx(args.input_pdf, args.intermediate_docx)

    print("Carregando tradutor offline…")
    translator = get_offline_translation_func(args.src, args.dest)

    print("Traduzindo DOCX offline…")
    translate_docx(args.intermediate_docx, args.output_docx, translator)

    print("Concluído! Revise", args.output_docx, "e exporte para PDF.")
if __name__ == "__main__":
    main()
